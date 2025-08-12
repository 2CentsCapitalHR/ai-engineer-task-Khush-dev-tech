import os
import json
import base64
import re
from io import BytesIO
from typing import List, Dict, Tuple
import streamlit as st
from dotenv import load_dotenv
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from PyPDF2 import PdfReader

from langchain_openai import AzureOpenAIEmbeddings, AzureChatOpenAI
from langchain.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQA

load_dotenv()

ADGM_CHECKLISTS = {
    "Company Incorporation": [
        "Articles of Association",
        "Memorandum of Association",
        "Incorporation Application Form",
        "UBO Declaration Form",
        "Register of Members and Directors"
    ]
}

DOC_TYPE_KEYWORDS = {
    "Articles of Association": ["articles of association", "aoa", "articles"],
    "Memorandum of Association": ["memorandum of association", "moa", "memorandum"],
    "Incorporation Application Form": ["incorporation application"],
    "UBO Declaration Form": ["ubo declaration"],
    "Register of Members and Directors": ["register of members", "register of directors"]
}

# --- Helper Functions ---
def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    doc = Document(BytesIO(file_bytes))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)

def classify_document(text: str) -> str:
    text_l = (text or "").lower()
    for doc_type, keys in DOC_TYPE_KEYWORDS.items():
        if any(k in text_l for k in keys):
            return doc_type
    return "Unknown Document Type"

def ingest_pdfs_to_faiss(pdf_file_objs: List[BytesIO]) -> Tuple[FAISS, str]:
    texts = []
    for pf in pdf_file_objs:
        pf.seek(0)
        reader = PdfReader(pf)
        pages = [page.extract_text() or "" for page in reader.pages]
        texts.append("\n".join(pages))
    chunks = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=150).split_text("\n".join(texts))
    embeddings = AzureOpenAIEmbeddings(
        deployment=os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT"),
        model="text-embedding-3-small",
        api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        azure_endpoint=os.getenv("AZURE_OPENAI_API_BASE"),
        api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
    )
    vectordb = FAISS.from_texts(chunks, embeddings)
    return vectordb, "ADGM Reference"

def _extract_json_from_text(text: str) -> Dict:
    start = text.find("{")
    if start == -1: return {}
    depth = 0
    for i in range(start, len(text)):
        if text[i] == "{": depth += 1
        elif text[i] == "}":
            depth -= 1
            if depth == 0:
                try:
                    return json.loads(text[start:i+1])
                except:
                    return {}
    return {}

def run_rag_review(retriever, doc_name: str, doc_text: str) -> Dict:
    llm = AzureChatOpenAI(
        deployment_name=os.getenv("AZURE_OPENAI_CHAT_DEPLOYMENT"),
        api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        azure_endpoint=os.getenv("AZURE_OPENAI_API_BASE"),
        api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
        temperature=0.0,
    )
    qa = RetrievalQA.from_chain_type(llm=llm, retriever=retriever)
    prompt = (
        "Review the document for ADGM compliance issues.\n"
        "Output JSON only with fields: document, issues:[{section, issue, severity, suggestion}].\n"
        f"Document name: {doc_name}\n\nContent:\n{doc_text}"
    )
    resp = qa.run(prompt)
    parsed = _extract_json_from_text(resp)
    if not parsed:
        parsed = {"document": doc_name, "issues": []}
    return parsed

def annotate_docx_inline(file_bytes: bytes, json_review: Dict) -> bytes:
    doc = Document(BytesIO(file_bytes))
    issues = json_review.get("issues", [])
    for issue in issues:
        section = (issue.get("section") or "").lower()
        snippet = section or issue.get("issue", "")[:50].lower()
        comment = f"[Suggestion: {issue.get('suggestion', '')}]"
        found = False
        for para in doc.paragraphs:
            if snippet and snippet in para.text.lower():
                run = para.add_run(" " + comment)
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                found = True
                break
        if not found:
            p = doc.add_paragraph(comment)
            p.runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()

# --- Streamlit App ---
st.set_page_config(page_title="ADGM Review", layout="wide")
st.title("🏛️ ADGM Compliance Review")

if "results" not in st.session_state:
    st.session_state.results = None

if not st.session_state.results:
    pdf_files = st.file_uploader("Upload ADGM Reference PDFs", accept_multiple_files=True, type=["pdf"])
    docx_files = st.file_uploader("Upload .docx files for review", accept_multiple_files=True, type=["docx"])
    if st.button("Run Review"):
        if not pdf_files or not docx_files:
            st.error("Please upload both reference PDFs and .docx files.")
        else:
            pdf_objs = [BytesIO(f.read()) for f in pdf_files]
            vectordb, _ = ingest_pdfs_to_faiss(pdf_objs)
            retriever = vectordb.as_retriever(search_kwargs={"k": 4})

            required_docs = ADGM_CHECKLISTS["Company Incorporation"]
            detected_types = []
            all_issues = []
            reviewed_files = []

            for docx in docx_files:
                raw = docx.read()
                text = extract_text_from_docx_bytes(raw)
                doc_type = classify_document(text)
                detected_types.append(doc_type)

                review_json = run_rag_review(retriever, doc_type, text)
                all_issues.extend([{
                    "document": doc_type,
                    "section": i.get("section", ""),
                    "issue": i.get("issue", ""),
                    "severity": i.get("severity", ""),
                    "suggestion": i.get("suggestion", "")
                } for i in review_json.get("issues", [])])

                reviewed_bytes = annotate_docx_inline(raw, review_json)
                reviewed_files.append((docx.name, reviewed_bytes))

            missing = [d for d in required_docs if d not in detected_types]
            output_json = {
                "process": "Company Incorporation",
                "documents_uploaded": len(docx_files),
                "required_documents": len(required_docs),
                "missing_document": missing[0] if missing else None,
                "issues_found": all_issues
            }

            st.session_state.results = {
                "json": output_json,
                "reviewed_files": reviewed_files
            }
            st.rerun()

else:
    st.subheader("Review JSON Output")
    st.json(st.session_state.results["json"])

    st.subheader("Reviewed Documents")
    for name, data in st.session_state.results["reviewed_files"]:
        st.download_button(
            label=f"Download reviewed {name}",
            data=data,
            file_name=f"reviewed_{name}",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    if st.button("🔄 Review New Documents"):
        st.session_state.results = None
        st.rerun()
